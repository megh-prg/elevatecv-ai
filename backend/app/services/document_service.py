import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen.canvas import Canvas


def generate_docx_resume(data: dict, title: str = "Optimized Resume") -> bytes:
    document = Document()
    document.add_heading(title, level=0)

    if summary := data.get("summary"):
        document.add_heading("Summary", level=1)
        document.add_paragraph(summary)

    if skills := data.get("skills"):
        document.add_heading("Skills", level=1)
        skills_paragraph = document.add_paragraph()
        skills_paragraph.add_run(", ".join(skills)).font.size = Pt(11)

    for section_name in ["experience", "projects"]:
        items = data.get(section_name) or []
        if not items:
            continue
        document.add_heading(section_name.capitalize(), level=1)
        for item in items:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(item)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_pdf_resume(data: dict, title: str = "Optimized Resume") -> bytes:
    buffer = io.BytesIO()
    canvas = Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 50
    y = height - margin
    canvas.setFont("Helvetica-Bold", 18)
    canvas.drawString(margin, y, title)
    y -= 30

    canvas.setFont("Helvetica", 11)

    def write_block(header: str, body: str) -> int:
        nonlocal y
        canvas.setFont("Helvetica-Bold", 12)
        canvas.drawString(margin, y, header)
        y -= 18
        canvas.setFont("Helvetica", 11)
        for line in body.splitlines():
            if y < margin + 30:
                canvas.showPage()
                y = height - margin
                canvas.setFont("Helvetica", 11)
            canvas.drawString(margin, y, line)
            y -= 14
        y -= 10
        return y

    if summary := data.get("summary"):
        y = write_block("Summary", summary)

    if skills := data.get("skills"):
        y = write_block("Skills", ", ".join(skills))

    for section_name in ["experience", "projects"]:
        items = data.get(section_name) or []
        if items:
            y = write_block(section_name.capitalize(), "\n".join(items))

    canvas.save()
    buffer.seek(0)
    return buffer.getvalue()
